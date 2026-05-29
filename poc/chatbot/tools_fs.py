"""Read-only file-system tools sandboxed under a docs root.

Every public function takes `docs_root: Path` so the module is pure and easy
to test. The dispatcher in `tools.py` binds the active `DOCS_ROOT` at call
time.

read_doc dispatches on file extension:
- text-like (.md, .txt, .py, etc.): UTF-8 decode
- .pdf: extract text via pypdf
- .docx: extract paragraphs + tables via python-docx
- image (.png, .jpg, etc.): OCR via pytesseract (requires the tesseract binary,
  `brew install tesseract tesseract-lang` on macOS)
"""
from __future__ import annotations

from pathlib import Path
from typing import Dict, List, Tuple

_MAX_FILES = 200
_MAX_BYTES = 32 * 1024
_SNIPPET_LEN = 160

# Treated as plain UTF-8 text.
_TEXT_EXTS = {
    "", ".md", ".markdown", ".txt", ".rst", ".log",
    ".csv", ".tsv", ".json", ".yaml", ".yml", ".toml", ".ini",
    ".py", ".js", ".ts", ".tsx", ".jsx", ".html", ".htm", ".css",
    ".sh", ".bash", ".zsh", ".sql", ".xml",
}
_PDF_EXTS = {".pdf"}
_DOCX_EXTS = {".docx"}
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".tif"}


def _resolve_path(docs_root: Path, rel_or_abs: str) -> Path:
    """Resolve a user-supplied path and reject anything outside docs_root.

    Raises ValueError if the resolved path escapes the sandbox.
    """
    root = docs_root.resolve()
    candidate = (root / rel_or_abs).resolve() if not Path(rel_or_abs).is_absolute() else Path(rel_or_abs).resolve()
    try:
        candidate.relative_to(root)
    except ValueError as e:
        raise ValueError(f"path '{rel_or_abs}' resolves outside DOCS_ROOT") from e
    return candidate


def list_docs(docs_root: Path, subdir: str = ".") -> List[Dict[str, object]]:
    """Recursively list files under docs_root/subdir.

    Returns at most _MAX_FILES entries, each {"path": str (relative to docs_root), "size_bytes": int}.
    Hidden files and directories (leading dot) are skipped.
    """
    root = docs_root.resolve()
    start = _resolve_path(docs_root, subdir)
    if not start.exists() or not start.is_dir():
        raise ValueError(f"subdir '{subdir}' is not a directory under DOCS_ROOT")
    out: List[Dict[str, object]] = []
    for p in sorted(start.rglob("*")):
        if any(part.startswith(".") for part in p.relative_to(root).parts):
            continue
        if not p.is_file():
            continue
        out.append({
            "path": p.relative_to(root).as_posix(),
            "size_bytes": p.stat().st_size,
        })
        if len(out) >= _MAX_FILES:
            break
    return out


def read_doc(docs_root: Path, path: str) -> Dict[str, object]:
    """Read a document under docs_root and return its text contents.

    Supports text/markdown, PDF (.pdf), Word (.docx), and images (OCR).
    Output is truncated to 32 KB. Image OCR requires the `tesseract` binary.
    """
    target = _resolve_path(docs_root, path)
    if not target.exists() or not target.is_file():
        raise FileNotFoundError(f"file '{path}' not found under DOCS_ROOT")

    ext = target.suffix.lower()
    rel = target.relative_to(docs_root.resolve()).as_posix()

    if ext in _TEXT_EXTS:
        text, truncated = _read_text(target)
        content_type = "text"
    elif ext in _PDF_EXTS:
        text, truncated = _read_pdf(target)
        content_type = "pdf"
    elif ext in _DOCX_EXTS:
        text, truncated = _read_docx(target)
        content_type = "docx"
    elif ext in _IMAGE_EXTS:
        text, truncated = _read_image(target)
        content_type = f"image-ocr ({ext.lstrip('.')})"
    else:
        raise ValueError(
            f"unsupported file extension '{ext or '(none)'}'. "
            f"Supported: text/markdown, .pdf, .docx, images (png/jpg/...)"
        )

    return {
        "path": rel,
        "content": text,
        "content_type": content_type,
        "truncated": truncated,
    }


def _truncate_text(text: str) -> Tuple[str, bool]:
    encoded = text.encode("utf-8")
    if len(encoded) <= _MAX_BYTES:
        return text, False
    return encoded[:_MAX_BYTES].decode("utf-8", errors="replace"), True


def _read_text(target: Path) -> Tuple[str, bool]:
    raw = target.read_bytes()
    truncated = len(raw) > _MAX_BYTES
    return raw[:_MAX_BYTES].decode("utf-8", errors="replace"), truncated


def _read_pdf(target: Path) -> Tuple[str, bool]:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError("pypdf is required for PDF reading. pip install pypdf") from e
    reader = PdfReader(str(target))
    parts = [(page.extract_text() or "") for page in reader.pages]
    return _truncate_text("\n\n".join(parts).strip())


def _read_docx(target: Path) -> Tuple[str, bool]:
    try:
        import docx
    except ImportError as e:
        raise RuntimeError("python-docx is required for .docx reading. pip install python-docx") from e
    document = docx.Document(str(target))
    parts = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return _truncate_text("\n".join(parts).strip())


def _read_image(target: Path) -> Tuple[str, bool]:
    try:
        import pytesseract
        from PIL import Image
    except ImportError as e:
        raise RuntimeError(
            "pytesseract and Pillow are required for image OCR. "
            "pip install pytesseract Pillow"
        ) from e
    img = Image.open(target)
    try:
        text = pytesseract.image_to_string(img, lang="eng+chi_sim")
    except pytesseract.TesseractNotFoundError as e:
        raise RuntimeError(
            "tesseract binary not found. On macOS: brew install tesseract tesseract-lang"
        ) from e
    except pytesseract.TesseractError:
        # Language pack (e.g. chi_sim) missing -- fall back to English only.
        text = pytesseract.image_to_string(img, lang="eng")
    return _truncate_text(text.strip())


def search_docs(docs_root: Path, query: str, max_hits: int = 20) -> List[Dict[str, object]]:
    """Case-insensitive substring search across files under docs_root.

    Returns at most max_hits entries of {"path", "line", "snippet"}.
    Skipped: hidden files/dirs, binaries (non-UTF-8), files larger than 1 MB.
    Note: only files readable as UTF-8 text are scanned. PDFs, DOCX, and
    images do not participate in keyword search -- use read_doc on them.
    """
    if not query:
        raise ValueError("query must be non-empty")
    needle = query.lower()
    root = docs_root.resolve()
    hits: List[Dict[str, object]] = []
    files_scanned = 0
    for p in sorted(root.rglob("*")):
        if any(part.startswith(".") for part in p.relative_to(root).parts):
            continue
        if not p.is_file():
            continue
        if p.stat().st_size > 1 * 1024 * 1024:
            continue
        files_scanned += 1
        if files_scanned > _MAX_FILES:
            break
        try:
            text = p.read_text(encoding="utf-8")
        except (UnicodeDecodeError, OSError):
            continue
        for line_no, line in enumerate(text.splitlines(), start=1):
            if needle in line.lower():
                col = line.lower().index(needle)
                start = max(0, col - 40)
                snippet = line[start:start + _SNIPPET_LEN]
                hits.append({
                    "path": p.relative_to(root).as_posix(),
                    "line": line_no,
                    "snippet": snippet,
                })
                if len(hits) >= max_hits:
                    return hits
    return hits
