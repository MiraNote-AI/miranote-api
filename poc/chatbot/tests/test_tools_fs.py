from __future__ import annotations
from pathlib import Path
import pytest

from poc.chatbot import tools_fs


def test_resolve_path_accepts_relative_inside_root(tmp_docs: Path):
    p = tools_fs._resolve_path(tmp_docs, "alpha.md")
    assert p == tmp_docs / "alpha.md"


def test_resolve_path_accepts_nested(tmp_docs: Path):
    p = tools_fs._resolve_path(tmp_docs, "nested/gamma.md")
    assert p == tmp_docs / "nested" / "gamma.md"


def test_resolve_path_rejects_parent_escape(tmp_docs: Path):
    with pytest.raises(ValueError, match="outside"):
        tools_fs._resolve_path(tmp_docs, "../escape.md")


def test_resolve_path_rejects_absolute_outside(tmp_docs: Path):
    with pytest.raises(ValueError, match="outside"):
        tools_fs._resolve_path(tmp_docs, "/etc/passwd")


def test_list_docs_root_returns_top_level(tmp_docs):
    out = tools_fs.list_docs(tmp_docs, ".")
    paths = sorted(item["path"] for item in out)
    assert paths == ["alpha.md", "beta.md", "nested/gamma.md"]
    for item in out:
        assert isinstance(item["size_bytes"], int) and item["size_bytes"] > 0


def test_list_docs_subdir(tmp_docs):
    out = tools_fs.list_docs(tmp_docs, "nested")
    paths = [item["path"] for item in out]
    assert paths == ["nested/gamma.md"]


def test_list_docs_rejects_escape(tmp_docs):
    with pytest.raises(ValueError, match="outside"):
        tools_fs.list_docs(tmp_docs, "../")


def test_read_doc_returns_content(tmp_docs):
    out = tools_fs.read_doc(tmp_docs, "alpha.md")
    assert out["path"] == "alpha.md"
    assert "First doc body." in out["content"]
    assert out["truncated"] is False


def test_read_doc_truncates_large_file(tmp_docs):
    big = tmp_docs / "big.md"
    big.write_text("x" * (40 * 1024), encoding="utf-8")  # 40 KB
    out = tools_fs.read_doc(tmp_docs, "big.md")
    assert len(out["content"].encode("utf-8")) <= 32 * 1024
    assert out["truncated"] is True


def test_read_doc_rejects_escape(tmp_docs):
    with pytest.raises(ValueError, match="outside"):
        tools_fs.read_doc(tmp_docs, "../escape.md")


def test_read_doc_missing_file(tmp_docs):
    with pytest.raises(FileNotFoundError):
        tools_fs.read_doc(tmp_docs, "nope.md")


def test_search_docs_finds_substring(tmp_docs):
    hits = tools_fs.search_docs(tmp_docs, "apples")
    paths = sorted(h["path"] for h in hits)
    assert paths == ["alpha.md", "nested/gamma.md"]
    for h in hits:
        assert "apples" in h["snippet"].lower()
        assert h["line"] >= 1


def test_search_docs_case_insensitive(tmp_docs):
    assert tools_fs.search_docs(tmp_docs, "BANANAS")[0]["path"] == "beta.md"


def test_search_docs_max_hits(tmp_docs):
    busy = tmp_docs / "busy.md"
    busy.write_text("\n".join(["needle"] * 50), encoding="utf-8")
    hits = tools_fs.search_docs(tmp_docs, "needle", max_hits=5)
    assert len(hits) == 5


def test_search_docs_snippet_truncated(tmp_docs):
    long = tmp_docs / "long.md"
    long.write_text("x" * 500 + "needle" + "y" * 500, encoding="utf-8")
    hits = tools_fs.search_docs(tmp_docs, "needle")
    assert len(hits) == 1
    assert len(hits[0]["snippet"]) <= 160
    assert "needle" in hits[0]["snippet"]


def test_search_docs_empty_query_rejected(tmp_docs):
    with pytest.raises(ValueError, match="query"):
        tools_fs.search_docs(tmp_docs, "")


# -- Multi-format read_doc --

def test_read_doc_text_reports_content_type(tmp_docs):
    out = tools_fs.read_doc(tmp_docs, "alpha.md")
    assert out["content_type"] == "text"


def test_read_doc_unsupported_extension(tmp_docs):
    (tmp_docs / "binary.exe").write_bytes(b"\x00\x01\x02")
    with pytest.raises(ValueError, match="unsupported file extension"):
        tools_fs.read_doc(tmp_docs, "binary.exe")


def test_read_doc_pdf(tmp_docs):
    pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    pdf_path = tmp_docs / "sample.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    c.drawString(72, 720, "MiraNote Architecture Snapshot")
    c.drawString(72, 700, "Section 1: Overview")
    c.drawString(72, 680, "The system uses a tool-calling chat loop.")
    c.save()

    out = tools_fs.read_doc(tmp_docs, "sample.pdf")
    assert out["content_type"] == "pdf"
    assert out["path"] == "sample.pdf"
    assert "Architecture Snapshot" in out["content"]
    assert "tool-calling" in out["content"]
    assert out["truncated"] is False


def test_read_doc_docx(tmp_docs):
    pytest.importorskip("docx")
    import docx
    docx_path = tmp_docs / "notes.docx"
    document = docx.Document()
    document.add_heading("Q2 Retro", level=1)
    document.add_paragraph("What went well: shipped two POCs.")
    document.add_paragraph("What to improve: spec discipline.")
    document.save(str(docx_path))

    out = tools_fs.read_doc(tmp_docs, "notes.docx")
    assert out["content_type"] == "docx"
    assert "Q2 Retro" in out["content"]
    assert "shipped two POCs" in out["content"]
    assert out["truncated"] is False


def _tesseract_available():
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


@pytest.mark.skipif(not _tesseract_available(), reason="tesseract binary not installed")
def test_read_doc_image_ocr(tmp_docs):
    from PIL import Image, ImageDraw, ImageFont
    img_path = tmp_docs / "note.png"
    img = Image.new("RGB", (600, 200), color="white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 32)
    except OSError:
        font = ImageFont.load_default()
    draw.text((20, 80), "HELLO MIRANOTE", fill="black", font=font)
    img.save(str(img_path))

    out = tools_fs.read_doc(tmp_docs, "note.png")
    assert out["content_type"].startswith("image-ocr")
    assert "MIRANOTE" in out["content"].upper()


def test_read_doc_image_without_tesseract_raises_friendly_error(tmp_docs, monkeypatch):
    # Simulate the binary being unavailable even if the library is installed.
    from PIL import Image
    img_path = tmp_docs / "stub.png"
    Image.new("RGB", (10, 10), color="white").save(str(img_path))

    import pytesseract

    def fake_image_to_string(*args, **kwargs):
        raise pytesseract.TesseractNotFoundError()

    monkeypatch.setattr(pytesseract, "image_to_string", fake_image_to_string)
    with pytest.raises(RuntimeError, match="brew install tesseract"):
        tools_fs.read_doc(tmp_docs, "stub.png")
